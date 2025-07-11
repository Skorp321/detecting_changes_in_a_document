import os
import logging
from typing import Dict, List, Optional, Tuple
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor
import chardet

# Document processing imports
from docx import Document
import PyPDF2
from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams
import re

from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document formats"""
    
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=4)
        self.supported_formats = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'txt': self._process_txt,
        }
    
    async def process_document(self, file_path: str) -> Dict:
        """Process document and extract text with metadata"""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Файл не найден: {file_path}")
            
            # Determine file format
            file_extension = Path(file_path).suffix.lower().lstrip('.')
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Неподдерживаемый формат файла: {file_extension}")
            
            # Process document in thread pool
            loop = asyncio.get_event_loop()
            processor = self.supported_formats[file_extension]
            
            result = await loop.run_in_executor(
                self.executor, 
                processor, 
                file_path
            )
            
            # Post-process text
            processed_text = self._post_process_text(result['text'])
            
            return {
                'text': processed_text,
                'metadata': result['metadata'],
                'paragraphs': self._split_into_paragraphs(processed_text),
                'word_count': len(processed_text.split()),
                'char_count': len(processed_text),
                'language': self._detect_language(processed_text)
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки документа {file_path}: {e}")
            raise
    
    def _process_pdf(self, file_path: str) -> Dict:
        """Process PDF document"""
        try:
            # Try PyPDF2 first
            text = ""
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'format': 'pdf',
                    'title': pdf_reader.metadata.title if pdf_reader.metadata else None,
                    'author': pdf_reader.metadata.author if pdf_reader.metadata else None,
                    'creator': pdf_reader.metadata.creator if pdf_reader.metadata else None,
                    'producer': pdf_reader.metadata.producer if pdf_reader.metadata else None,
                }
                
                # Extract text
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Ошибка извлечения текста со страницы {page_num}: {e}")
            
            # If PyPDF2 didn't extract much text, try pdfminer
            if len(text.strip()) < 100:
                try:
                    text = extract_text(
                        file_path,
                        laparams=LAParams(
                            boxes_flow=0.5,
                            word_margin=0.1,
                            char_margin=2.0,
                            line_margin=0.5
                        )
                    )
                except Exception as e:
                    logger.warning(f"Ошибка извлечения текста с помощью pdfminer: {e}")
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки PDF: {e}")
            raise
    
    def _process_docx(self, file_path: str) -> Dict:
        """Process DOCX document"""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Extract metadata
            metadata = {
                'format': 'docx',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'title': doc.core_properties.title,
                'author': doc.core_properties.author,
                'created': doc.core_properties.created,
                'modified': doc.core_properties.modified,
            }
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки DOCX: {e}")
            raise
    
    def _process_txt(self, file_path: str) -> Dict:
        """Process TXT document"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding'] or 'utf-8'
            
            # Read text file
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
            
            # Get file metadata
            file_stats = os.stat(file_path)
            metadata = {
                'format': 'txt',
                'encoding': encoding,
                'size': file_stats.st_size,
                'lines': len(text.splitlines()),
            }
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки TXT: {e}")
            raise
    
    def _post_process_text(self, text: str) -> str:
        """Post-process extracted text"""
        if not text:
            return ""
        
        # Normalize line breaks first (convert different line endings)
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive line breaks (more than 2 consecutive)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove trailing spaces from lines but preserve line structure
        lines = text.split('\n')
        lines = [line.rstrip() for line in lines]
        text = '\n'.join(lines)
        
        # Remove excessive spaces within lines (but preserve single spaces and line breaks)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Fix common OCR errors for Russian text
        text = self._fix_russian_text(text)
        
        # Remove page headers/footers patterns
        text = self._remove_headers_footers(text)
        
        return text.strip()
    
    def _fix_russian_text(self, text: str) -> str:
        """Fix common Russian text issues"""
        # Common OCR substitutions
        replacements = {
            'о': 'o',  # Cyrillic o to Latin o
            'а': 'a',  # Cyrillic a to Latin a
            'е': 'e',  # Cyrillic e to Latin e
            'р': 'p',  # Cyrillic p to Latin p
            'у': 'y',  # Cyrillic y to Latin y
            'х': 'x',  # Cyrillic x to Latin x
        }
        
        # Apply replacements carefully (only if surrounded by other Cyrillic)
        for cyrillic, latin in replacements.items():
            # This is a simplified approach - in production would need more sophisticated logic
            pass
        
        return text
    
    def _remove_headers_footers(self, text: str) -> str:
        """Remove common header/footer patterns"""
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            
            # Skip lines that look like headers/footers
            if self._is_header_footer(line):
                continue
                
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _is_header_footer(self, line: str) -> bool:
        """Check if line is likely a header or footer"""
        line = line.strip()
        
        # Empty lines
        if not line:
            return False
        
        # Page numbers
        if re.match(r'^\d+$', line):
            return True
        
        # Common header/footer patterns
        header_patterns = [
            r'стр\.\s*\d+',  # page numbers
            r'страница\s*\d+',
            r'^\d+\s*$',  # standalone numbers
            r'^\d+\s*/\s*\d+$',  # page x/y
        ]
        
        for pattern in header_patterns:
            if re.search(pattern, line.lower()):
                return True
        
        return False
    
    def _split_into_paragraphs(self, text: str) -> List[str]:
        """Split text into logical paragraphs"""
        if not text:
            return []
        
        # Try multiple splitting strategies
        paragraphs = []
        
        # Strategy 1: Split by double newlines
        double_newline_split = text.split('\n\n')
        if len(double_newline_split) > 1:
            paragraphs.extend(double_newline_split)
        else:
            # Strategy 2: Split by single newlines if no double newlines
            single_newline_split = text.split('\n')
            if len(single_newline_split) > 1:
                paragraphs.extend(single_newline_split)
            else:
                # Strategy 3: Split by sentences if no line breaks
                sentence_split = re.split(r'[.!?]+\s+', text)
                paragraphs.extend(sentence_split)
        
        # Clean up paragraphs
        cleaned_paragraphs = []
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if paragraph and len(paragraph) > 5:  # More lenient minimum length
                # Normalize whitespace within paragraphs but keep structure
                paragraph = re.sub(r'\s+', ' ', paragraph)
                cleaned_paragraphs.append(paragraph)
        
        # If we still have no paragraphs, treat the whole text as one paragraph
        if not cleaned_paragraphs and text.strip():
            cleaned_paragraphs = [text.strip()]
        
        return cleaned_paragraphs
    
    def _detect_language(self, text: str) -> str:
        """Detect document language"""
        if not text:
            return 'unknown'
        
        # Simple Russian text detection
        cyrillic_chars = len(re.findall(r'[а-яё]', text.lower()))
        latin_chars = len(re.findall(r'[a-z]', text.lower()))
        
        if cyrillic_chars > latin_chars:
            return 'ru'
        elif latin_chars > cyrillic_chars:
            return 'en'
        else:
            return 'mixed'
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats"""
        return list(self.supported_formats.keys())
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """Validate if file can be processed"""
        try:
            if not os.path.exists(file_path):
                return False, "Файл не найден"
            
            file_extension = Path(file_path).suffix.lower().lstrip('.')
            if file_extension not in settings.ALLOWED_EXTENSIONS_LIST:
                return False, f"Неподдерживаемый формат: {file_extension}"
            
            file_size = os.path.getsize(file_path)
            if file_size > settings.MAX_FILE_SIZE:
                return False, f"Файл слишком большой: {file_size} байт"
            
            return True, "OK"
            
        except Exception as e:
            return False, f"Ошибка валидации: {e}" 