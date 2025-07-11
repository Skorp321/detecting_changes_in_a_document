import logging
import csv
import io
from typing import List, Dict, Any
from datetime import datetime

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Service for generating reports from analysis results"""
    
    def __init__(self):
        pass
    
    async def generate_report(self, results: List[Dict], format: str) -> io.BytesIO:
        """Generate report in specified format"""
        try:
            if format == "csv":
                return await self._generate_csv(results)
            elif format == "pdf":
                return await self._generate_pdf(results)
            elif format == "word":
                return await self._generate_word(results)
            else:
                raise ValueError(f"Unsupported format: {format}")
                
        except Exception as e:
            logger.error(f"Error generating report: {e}")
            raise
    
    async def _generate_csv(self, results: List[Dict]) -> io.BytesIO:
        """Generate CSV report"""
        output = io.BytesIO()
        
        # Create string buffer for CSV
        csv_buffer = io.StringIO()
        
        fieldnames = [
            'Оригинальный текст',
            'Измененный текст', 
            'Комментарии LLM',
            'Необходимые согласования',
            'Тип изменения',
            'Критичность',
            'Уверенность',
            'Дата создания'
        ]
        
        writer = csv.DictWriter(csv_buffer, fieldnames=fieldnames)
        writer.writeheader()
        
        for result in results:
            writer.writerow({
                'Оригинальный текст': result.get('originalText', ''),
                'Измененный текст': result.get('modifiedText', ''),
                'Комментарии LLM': result.get('llmComment', ''),
                'Необходимые согласования': ', '.join(result.get('requiredServices', [])),
                'Тип изменения': result.get('changeType', ''),
                'Критичность': result.get('severity', ''),
                'Уверенность': result.get('confidence', ''),
                'Дата создания': result.get('createdAt', '')
            })
        
        # Convert to bytes
        csv_content = csv_buffer.getvalue()
        output.write(csv_content.encode('utf-8-sig'))  # BOM for Excel
        output.seek(0)
        
        return output
    
    async def _generate_pdf(self, results: List[Dict]) -> io.BytesIO:
        """Generate PDF report - placeholder implementation"""
        # For now, return a simple text-based PDF placeholder
        output = io.BytesIO()
        
        content = f"""ОТЧЕТ ОБ АНАЛИЗЕ ИЗМЕНЕНИЙ В ДОКУМЕНТАХ

Дата генерации: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}
Количество изменений: {len(results)}

{"="*60}

"""
        
        for i, result in enumerate(results, 1):
            content += f"""
ИЗМЕНЕНИЕ {i}
{"-"*20}
Оригинальный текст: {result.get('originalText', '')[:100]}...
Измененный текст: {result.get('modifiedText', '')[:100]}...
Комментарий: {result.get('llmComment', '')}
Необходимые согласования: {', '.join(result.get('requiredServices', []))}
Критичность: {result.get('severity', '')}

"""
        
        output.write(content.encode('utf-8'))
        output.seek(0)
        
        return output
    
    async def _generate_word(self, results: List[Dict]) -> io.BytesIO:
        """Generate Word document report"""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Акт разногласий', 0)
            title.alignment = 1  # Center alignment
            
            # Metadata
            doc.add_paragraph(f'Дата генерации: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
            doc.add_paragraph(f'Количество изменений: {len(results)}')
            doc.add_paragraph('')
            
            # Results table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Оригинальный текст'
            hdr_cells[1].text = 'Измененный текст'
            hdr_cells[2].text = 'Комментарии'
            hdr_cells[3].text = 'Согласования'
            
            # Data rows
            for result in results:
                row_cells = table.add_row().cells
                row_cells[0].text = result.get('originalText', '')[:200] + ('...' if len(result.get('originalText', '')) > 200 else '')
                row_cells[1].text = result.get('modifiedText', '')[:200] + ('...' if len(result.get('modifiedText', '')) > 200 else '')
                row_cells[2].text = result.get('llmComment', '')
                row_cells[3].text = ', '.join(result.get('requiredServices', []))
            
            # Save to BytesIO
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            return output
            
        except ImportError:
            # Fallback to text if python-docx is not available
            logger.warning("python-docx not available, generating text file")
            return await self._generate_pdf(results) 